import os
import shutil
import warnings
from io import BytesIO
from pathlib import Path
import pathlib
from typing import List

import requests
from docx import Document as WordDocument
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from PIL import Image as PILimage
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait

warnings.filterwarnings('ignore')
os.chdir(os.path.dirname(os.path.abspath(__file__)))

designerToWeb = {
    "1 Switch Plate": "1 or 2 Switch Plate",
    "2 Switch Plate": "1 or 2 Switch Plate",
    "3 Switch Plate": "3 Switch Plate",
    "4 Switch Plate": "4 Switch Plate",
    "6 Switch Plate": "6 Switch Plate",
    "8 Switch Plate - HR": "8 Switch - Horizontal",
    "8 Switch Plate - VR": "8 Switch - Vertical",
    "12 Switch Plate": "12 Switch"
}


def xlToWebDict(sheet):
    if sheet == "Infinity":
        switch_types = ['1 Gang', '2 Gang', '3 Gang', '4 Gang', '1 Gang Profile Keypad', '2 Gang Profile Keypad', '3 Gang Profile Keypad', '4 Gang Profile Keypad', '6 Gang Profile Keypad', 'Blinds', '2 Blinds', 'Curtain',
                        '2 Curtain', 'Door Bell', 'Fan Dimmer', 'Light Dimmer', '2 Light Dimmer', '3 Light Dimmer', 'Tunable', 'Socket (5-15 Amps)', 'Socket (2 USB+Switch)', 'C Type', 'Socket with C Type', 'HDMI USB', 'Cable', 'Data', 'Telephone']
        XL_TO_WEB = {item: item for item in switch_types}
        XL_TO_WEB["2 Gang - WR(S)"] = '2 Gang'
        XL_TO_WEB["2 Gang (M)"] = '2 Gang'
        XL_TO_WEB["1 Gang (M)"] = '1 Gang'
        XL_TO_WEB["Socket (USB+C-type(2A)+Switch)"] = "Socket (2 USB+Switch)"
        XL_TO_WEB["Telephone Socket"] = "Telephone"

    else:
        switch_types = ['1 Gang', '2 Gang', '3 Gang', '4 Gang', 'Blinds', '2 Blinds', 'Curtain',
                        '2 Curtain', 'Door Bell', 'Fan Dimmer', 'Light Dimmer', '2 Light Dimmer', '3 Light Dimmer', 'Tunable', 'Socket (5-15 Amps)', 'Socket (2 USB+Switch)', 'C Type', 'Socket with C Type', 'HDMI USB', 'Cable', 'Data', 'Telephone']
        XL_TO_WEB = {item: item for item in switch_types}
        XL_TO_WEB["2 Gang - WR(S)"] = '2 Gang'
        XL_TO_WEB["1 Gang - WR(S)"] = "1 Gang"
        XL_TO_WEB["2 Gang (M)"] = '2 Gang'
        XL_TO_WEB["1 Gang (M)"] = '1 Gang'
        XL_TO_WEB["Socket (USB+C-type(2A)+Switch)"] = "Socket (2 USB+Switch)"
        # XL_TO_WEB["Telephone Socket"] = "Telephone"

        XL_TO_WEB["1 Gang Profile Keypad"] = "1 Gang"
        XL_TO_WEB["2 Gang Profile Keypad"] = "2 Gang"
        XL_TO_WEB["3 Gang Profile Keypad"] = "3 Gang"
        XL_TO_WEB["4 Gang Profile Keypad"] = "4 Gang"
        XL_TO_WEB["6 Gang Profile Keypad"] = "4 Gang"
        XL_TO_WEB["Dummy + Backbox"] = "1 Gang"
        XL_TO_WEB["Telephone Socket"] = "1 Gang"
        XL_TO_WEB["Foot Lamp"] = "1 Gang"

    return XL_TO_WEB


def findExcelFile():
    for file in os.listdir():
        try:
            wb = load_workbook(file, read_only=True)
        except InvalidFileException:
            continue
        else:
            return file
    raise Exception("No File Readable to Excel in this Directory")


def removeFiles(path, dir):
    print(os.listdir())
    if path in os.listdir():
        os.remove(path)

    if dir in os.listdir():
        shutil.rmtree(dir)


def setImageDpi(path, dpi):
    image = PILimage.open(path)
    image.convert("RGB").save(path, dpi=(dpi, dpi))


class Doc():
    def __init__(self, template="assets/template.docx", fileName="Proposal", logo="assets/logo.png"):
        self.doc = WordDocument(template)
        self.fileName = fileName
        self.logo = logo

    def addHeader(self):
        header = self.doc.sections[0].header
        paragraph = header.paragraphs[0]

        logo_run = paragraph.add_run()
        logo_run.add_picture(self.logo, width=Inches(2.5))

    def addRun(self, img, desc1, desc2,img2=None):
        p = self.doc.add_paragraph()
        r = p.add_run()
        if img2 != None:
            r.add_picture(img2)
        r.add_picture(img)
        r.add_break()
        r.add_text(desc1)
        r.add_break()
        r.add_text(desc2)

    def save(self):
        self.doc.save(f"{self.fileName}.docx")


def find(it, pred, default=None):
    return next(filter(pred, it), default)


class Sheet():
    # Goal for this sheet is to have all the info we want to extract
    # Column info for space and product descriptions
    # Column info for switch info
    # Potentially more column info for proposal information
    def __init__(self, name, num):
        self.name = name
        self.info = {}
        self.num = num
        self.maxRow = None

    def __repr__(self):
        return f"Sheet({self.name})"

    def addColInfo(self, colStart, info, colEnd=None):
        if colEnd == None:
            colEnd = colStart

        self.info[info] = [
            colStart, colEnd] if colStart != colEnd else colStart


class Agent():

    def __init__(self, wb, dir="tmp", sheets: List[Sheet] = []):
        self.options = webdriver.ChromeOptions()
        self.options.add_argument("--disable-gpu")
        self.options.add_argument("--headless")
        self.options.add_argument("--window-size=1920,1080")
        self.options.add_argument('log-level=3')
        self.driver = webdriver.Chrome(options=self.options)
        self.maxWait = 30
        self.wait = WebDriverWait(self.driver, self.maxWait)

        self.wb = wb
        self.dir = dir
        self.XL_WEB_INF = xlToWebDict("Infinity")
        self.XL_WEB_DES = xlToWebDict("Designer")
        self.startInd = 14

        self.sheetObjs = sheets
        self.sheets = [self.wb[sheet.name] for sheet in self.sheetObjs]
        self.maxRows = [self.getMaxLen(sheet) for sheet in self.sheets]

        self.indArr = [elem if self.maxRows[elem] >=  # Finds any empty sheets and marks them
                       self.startInd else None for elem in range(len(self.sheets))]

        def reformat(x): return [x[elem] for elem in range(
            len(x)) if elem == self.indArr[elem]]

        self.maxRows = reformat(self.maxRows)
        self.sheets = reformat(self.sheets)
        self.sheetObjs = reformat(self.sheetObjs)

        if self.dir in os.listdir():
            os.chdir(self.dir)
        else:
            os.mkdir(self.dir)
            os.chdir(self.dir)

    def openToIndia(self):
        # Waits until page is fully loaded
        self.driver.implicitly_wait(self.maxWait)
        self.driver.get("https://app.smarttouchswitch.com/")

        try:
            button = self.driver.find_element(
                By.CLASS_NAME, "build-action-button")
            button.click()  # Opens to the Builder Menu

        except TimeoutException:
            print("Loading took too much time!")

    def getMaxLen(self, sheet):
        # For Infinity and Designer B is the Date
        # For Infinity O is a Color
        # For Designer O is the System
        # These tell how many rows of data there is

        for i in range(self.startInd, sheet.max_row):
            if sheet[f"B{i}"].value == None and sheet[f"O{i}"].value == None:
                return i-1

    def getModules(self):
        self.modules = []
        for sheet in range(len(self.sheets)):

            start, end = self.sheetObjs[sheet].info["Modules"]
            cell_range = f"{start}{self.startInd}:{end}{self.maxRows[sheet]}"
            maxModuleSize = ord(end) - ord(start) + 1
            cnt = 0
            for column in self.sheets[sheet][cell_range]:
                for cell in column:
                    if cnt % maxModuleSize == 0:
                        self.modules.append([self.sheetObjs[sheet], cell.row])
                    if cell.value != None:
                        self.modules[cnt//maxModuleSize].append(cell.value)
                    cnt += 1
        tmpModules = []
        cnt = 0
        for module in self.modules:
            tmpModules.append([module[:2]])
            for item in module[2:]:
                XL_WEB = self.XL_WEB_INF if module[0].name == "Infinity" else self.XL_WEB_DES
                if item in XL_WEB:
                    tmpModules[cnt].append(XL_WEB[item])
                else:
                    continue
            # if a module is empty (has no valid switches or excel row is empty)
            if len(tmpModules[cnt]) == 1:
                tmpModules.pop()
            else:
                cnt += 1
        """
        self.modules : list[module] ->  
        module = [[Sheet(Infinity or Designer),rowOfInformation], Module 1, Module 2, Module 3]
        """
        self.modules = tmpModules

    def getColors(self):  # sheet is the indice of the sheet
        self.colors = []
        for sheet in range(len(self.sheets)):
            if self.sheetObjs[sheet].name == "Designer":  # Only for Designer currently
                start, end = self.sheetObjs[sheet].info["Colors"]
                cell_range = f"{start}{self.startInd}:{end}{self.maxRows[sheet]}"

                colorArr = ord(end) - ord(start) + 1
                cnt = 0
                for column in self.sheets[sheet][cell_range]:
                    for cell in column:
                        if cnt % colorArr == 0:
                            self.colors.append(
                                [[self.sheetObjs[sheet], cell.row]])
                        if cell.value != None:
                            self.colors[cnt//colorArr].append(cell.value)
                        cnt += 1
        # print(self.colors)

    def clickColor(self, level: str, colorProfile: str):
        # print(colorProfile)
        self.wait.until(
            EC.visibility_of_any_elements_located((By.CLASS_NAME, "mod-label"))
        )

        #colorProfile = [[Sheet(Designer), row], OuterGlass, OuterFrame, InnerGlass, InnerFrame]
        colorType = self.driver.find_element(
            By.XPATH, f"//span[text()=\'{level}\']"
        )
        colorType.click()
        if level == "Outer Surface":
            self.wait.until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "fab-label-inner"))
            )
            glass = self.driver.find_element(
                By.CLASS_NAME, "fab-label-inner"
            )
            glass.click()
        color = self.driver.find_element(
            By.XPATH, f"//a[@glass-color=\'{colorProfile.lower()}\'][@data-colortype='glass']"
        ).click()
    def clickModules(self):
        for moduleInd in range(len(self.modules)):
            modules = self.modules[moduleInd]
            self.openToIndia()

            self.wait.until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "mod-label"))
            )

            if modules[0][0].name == "Designer":
                ws = self.wb["Designer"]
                switchPlate = designerToWeb[ws[f"{modules[0][0].info['Switch']}{modules[0][1]}"].value]
                modType = self.driver.find_element(
                    By.XPATH, f"//span[text()=\'{switchPlate}\']")
                modType.click()
            else:
                # TODO: Implement vertical switches (when excel is updated)
                modType = self.driver.find_elements(
                    By.CLASS_NAME, "mod-label")[len(modules)-2]
                modType.click()

            modPanel = self.wait.until(
                lambda driver: driver.find_element(
                    By.CSS_SELECTOR, 'div[data-panelid=".modulePanel"]'
                )
            )
            modPanel.click()

            self.wait.until(
                EC.visibility_of_any_elements_located(
                    (By.CLASS_NAME, "module-type-label"))
            )

            # Skip the module info and get the modules themselves
            for module in modules[1:]:
                self.driver.find_element(
                    By.XPATH, f"//div[text()=\'{module}\']"
                ).click()

            modPanel.click()
            self.wait.until(
                EC.invisibility_of_element(
                    (By.CLASS_NAME, "module-type-label"))
            )
            # self.wait = WebDriverWait(self.driver,30)
            # #TODO: Implement Colors
            if modules[0][0].name == "Designer":
                colorPanel = self.driver.find_element(
                    By.CSS_SELECTOR, 'div[data-panelid=".colorPanel"]'
                )
                colorPanel.click()
                # self.driver.execute_script("arguments[0].click();", colorPanel)
                colorProfile = find(
                    self.colors, lambda x: x[0][1] == modules[0][1]
                )[1:]
                self.clickColor("Outer Surface", colorProfile=colorProfile[0])
                self.clickColor("Outer Frame", colorProfile=colorProfile[1])
                self.clickColor("Inner Surface", colorProfile=colorProfile[2])
                self.clickColor("Inner Frame", colorProfile=colorProfile[3])
                response = requests.get(f"https://app.smarttouchswitch.com/modules/components/images/frames/{colorProfile[0]}{colorProfile[1]}-Frame.png")
                frame = PILimage.open(BytesIO(response.content))
                framePath = f"frame_{moduleInd}.png"
                frame.save(framePath)
                modules[0].append(Path(framePath))
                colorPanel.click()


            self.screenshot(index=moduleInd)

    def screenshot(self, index):
        final_switch = self.driver.find_element(
            By.CLASS_NAME, "content-center")
        final_switch.screenshot(f'switch_{index}.png')
        print(f"Switch {index} Completed")

    def close(self):
        self.driver.quit()
        os.chdir("..")

    def getCol(self, sheet, colLetter):
        colVals = []
        for column in self.wb[self.sheetObjs[sheet].name][f"{colLetter}{self.startInd}:{colLetter}{self.maxRows[sheet]}"]:
            for cell in column:
                colVals.append(cell.value)
        return colVals

    def publish(self, fileName="Proposal", debug=False):
        document = Doc(fileName=fileName)
        document.addHeader()
        if debug:
            for img_path in os.listdir(self.dir):
                path = str(Path(os.path.join(self.dir, img_path)).absolute())
                spaceText = "Space: "
                prodText = "Product Description: "
                setImageDpi(path, 96*2)

                document.addRun(path, spaceText, prodText)
            document.save()
        else:
            for switch in range(len(self.modules)):
                sheetObj = self.modules[switch][0][0]
                row = self.modules[switch][0][1]
                frameImg = None
                if isinstance(self.modules[switch][0][-1],pathlib.PurePath):
                    frameImg = Path(os.path.join(self.dir, self.modules[switch][0][-1].absolute()))
                sheet = self.wb[sheetObj.name]

                prodDesc = "Product Description: " + \
                    str(sheet[f"{sheetObj.info['Product']}{row}"].value)
                space = "Space: " + \
                    str(sheet[f"{sheetObj.info['Space']}{row}"].value)
                path = str(
                    Path(os.path.join(self.dir, f"switch_{switch}.png")).absolute())
                setImageDpi(path, 96*2)
                document.addRun(path, space, prodDesc,frameImg)
            document.save()


dir = "tmp"
xlPath = findExcelFile()
removeFiles(Path(xlPath).with_suffix(".docx"), dir)
print(f"Excel File Found: {xlPath}")

wb = load_workbook(xlPath,
                   read_only=True, data_only=True)

infinity = Sheet("Infinity", 0)
infinity.addColInfo(info="Modules", colStart="K", colEnd="M")
infinity.addColInfo(info="Product", colStart="U")
infinity.addColInfo(info="Space", colStart="D")

designer = Sheet("Designer", 1)
designer.addColInfo(info="Modules", colStart="I", colEnd="N")
designer.addColInfo(info="Switch", colStart="G")
designer.addColInfo(info="Product", colStart="Y")
designer.addColInfo(info="Space", colStart="D")
designer.addColInfo(info="Colors", colStart="Q", colEnd="T")

agent = Agent(wb, dir=dir, sheets=[
    infinity,
    designer
])

agent.getModules()
agent.getColors()
agent.clickModules()
agent.close()

agent.publish(Path(xlPath).stem)
